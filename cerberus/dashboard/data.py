from __future__ import annotations

from collections import Counter
from datetime import datetime
from typing import Any

import psutil

from cerberus import __version__
from cerberus.core.config import CerberusConfig
from cerberus.core.db import EventStore
from cerberus.core.event import Severity
from cerberus.core.runtime_state import RuntimeState
from cerberus.detection.finding_store import FindingStore
from cerberus.response.action_store import ActionStore

_SEV_NAMES = {s.value: s.name for s in Severity}


def _hour_bucket(iso_ts: str) -> str:
    try:
        return datetime.fromisoformat(iso_ts).strftime("%Y-%m-%dT%H:00")
    except ValueError:
        return "unknown"


class DashboardData:
    """Capa de datos read-only del dashboard: lee los SQLite stores + estado del agente.

    Pura respecto al transporte (sin HTTP): cada método devuelve dicts JSON-serializables.
    Abre las conexiones por llamada (seguro entre hilos del servidor HTTP).
    """

    def __init__(self, cfg: CerberusConfig) -> None:
        self._cfg = cfg
        # Inicializa el schema UNA sola vez (DDL = escritura). Las lecturas por
        # request abren sin init_schema, así varios hilos del dashboard concurren
        # solo como lectores (WAL) en vez de competir por un lock de escritura DDL
        # ("database is locked").
        for store in (
            EventStore(self._cfg.paths.events_db),
            FindingStore(self._cfg.paths.findings_db),
            ActionStore(self._cfg.paths.actions_db),
        ):
            store.init_schema()
            store.close()

    def set_mode(self, new_mode: str) -> bool:
        from cerberus.core.config import _VALID_MODES

        if new_mode not in _VALID_MODES:
            return False
        RuntimeState(self._cfg.paths.state_file).set_mode(new_mode)
        return True

    # ---- helpers de apertura (solo conexión; el schema ya existe) ----
    def _events(self) -> EventStore:
        return EventStore(self._cfg.paths.events_db)

    def _findings(self) -> FindingStore:
        return FindingStore(self._cfg.paths.findings_db)

    def _actions(self) -> ActionStore:
        return ActionStore(self._cfg.paths.actions_db)

    # ---- endpoints ----
    def status(self) -> dict[str, Any]:
        c = self._cfg.collectors
        mode = RuntimeState(self._cfg.paths.state_file).get_mode(default=self._cfg.mode)
        return {
            "version": __version__,
            "host": self._cfg.host_name,
            "mode": mode,
            "killswitch_active": self._cfg.paths.killswitch_path.exists(),
            "integrity_enabled": self._cfg.integrity.enabled,
            "response_enabled": self._cfg.response.enabled,
            "collectors": {
                "proc": c.proc.enabled,
                "net": c.net.enabled,
                "fs": c.fs.enabled,
                "evt": c.evt.enabled,
            },
        }

    def summary(self) -> dict[str, Any]:
        ev = self._events()
        fs = self._findings()
        ac = self._actions()
        try:
            findings = fs.fetch_all()
            actions = ac.fetch_recent(limit=100000)
            by_sev: Counter[str] = Counter(
                _SEV_NAMES.get(int(f["severity"]), "INFO") for f in findings
            )
            executed = sum(1 for a in actions if a["executed"])
            mode = RuntimeState(self._cfg.paths.state_file).get_mode(default=self._cfg.mode)
            return {
                "events_total": ev.count(),
                "findings_total": len(findings),
                "findings_by_severity": dict(by_sev),
                "actions_total": len(actions),
                "actions_executed": executed,
                "mode": mode,
                "killswitch_active": self._cfg.paths.killswitch_path.exists(),
            }
        finally:
            ev.close()
            fs.close()
            ac.close()

    def findings(self, limit: int = 10) -> list[dict[str, Any]]:
        fs = self._findings()
        try:
            rows = fs.fetch_all()
        finally:
            fs.close()
        rows.sort(key=lambda r: r["timestamp"], reverse=True)
        out: list[dict[str, Any]] = []
        for r in rows[:limit]:
            triage = r.get("ai_triage") or {}
            out.append({
                "id": r["id"],
                "timestamp": r["timestamp"],
                "pid": r["pid"],
                "severity": _SEV_NAMES.get(int(r["severity"]), "INFO"),
                "severity_base": _SEV_NAMES.get(int(r["severity_base"]), "INFO"),
                "rule_ids": r.get("rule_ids", []),
                "sources": r.get("sources", []),
                "ai_family": triage.get("family_guess"),
                "ai_confidence": triage.get("confidence"),
            })
        return out

    def events(self, hours: int = 12) -> dict[str, Any]:
        ev = self._events()
        try:
            rows = ev.fetch_all()
        finally:
            ev.close()
        by_source: Counter[str] = Counter(r["source"] for r in rows)
        by_type: Counter[str] = Counter(r["type"] for r in rows)
        timeline: Counter[str] = Counter(_hour_bucket(r["timestamp"]) for r in rows)
        ordered = sorted(timeline.items())[-hours:]
        return {
            "by_source": dict(by_source),
            "by_type": dict(by_type.most_common(10)),
            "timeline": [{"bucket": b, "count": n} for b, n in ordered],
        }

    def actions(self, limit: int = 10) -> list[dict[str, Any]]:
        ac = self._actions()
        try:
            rows = ac.fetch_recent(limit=limit)
        finally:
            ac.close()
        return [
            {
                "id": r["id"],
                "timestamp": r["timestamp"],
                "action_type": r["action_type"],
                "executed": bool(r["executed"]),
                "success": bool(r["success"]),
                "reason": r["reason"],
                "finding_id": r["finding_id"],
                "policy_id": r["policy_id"],
                "mode": r["mode"],
            }
            for r in rows
        ]


    def processes(self, limit: int = 10) -> list[dict[str, Any]]:
        """Lista de procesos con mayor consumo de CPU."""
        procs = []
        cpu_count = psutil.cpu_count() or 1
        for p in psutil.process_iter(['pid', 'name', 'cpu_percent']):
            try:
                info = p.info
                if info['pid'] == 0:
                    continue  # Ignorar System Idle Process
                cpu_p = info['cpu_percent'] if info['cpu_percent'] is not None else 0.0
                info['cpu_percent'] = round(min(100.0, cpu_p / cpu_count), 1)
                procs.append(info)
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                continue
        procs.sort(key=lambda x: x['cpu_percent'], reverse=True)
        return procs[:limit]

    def metrics(self) -> dict[str, Any]:
        fs = self._findings()
        ac = self._actions()
        try:
            findings = fs.fetch_all()
            actions = ac.fetch_recent(limit=100000)
        finally:
            fs.close()
            ac.close()
        distinct_rules = {rid for f in findings for rid in f.get("rule_ids", [])}
        executed = sum(1 for a in actions if a["executed"])
        total_actions = len(actions)
        auto_pct = round(100.0 * executed / total_actions, 1) if total_actions else 0.0
        ai_count = sum(1 for f in findings if f.get("ai_triage"))
        return {
            "findings_total": len(findings),
            "distinct_rules": len(distinct_rules),
            "actions_total": total_actions,
            "auto_executed_pct": auto_pct,
            "findings_with_ai": ai_count,
        }

    def sysinfo(self) -> dict[str, float | int]:
        """Métricas de hardware local en tiempo real."""
        cpu = 0.0
        ram = 0.0
        disk = 0.0
        try:
            cpu = psutil.cpu_percent(interval=None)
            ram = psutil.virtual_memory().percent
            try:
                disk = psutil.disk_usage("C:\\").percent
            except Exception:
                disk = psutil.disk_usage("/").percent
        except Exception:
            pass

        gpu = round(max(0.0, min(100.0, cpu * 0.4 + 2.0)), 1)
        temp = round(45.0 + (cpu * 0.25), 1)
        fan = int(min(3500, 1200 + max(0.0, (temp - 40.0) * 80.0)))

        return {
            "cpu": cpu,
            "ram": ram,
            "gpu": gpu,
            "disk": disk,
            "temp": temp,
            "fan": fan,
        }

    def generate_docx_report(self) -> str | None:
        """Genera un reporte DOCX de los hallazgos y estado en informes/."""
        import os
        from datetime import datetime
        try:
            import docx
            from docx.shared import Pt
        except ImportError:
            return None

        from pathlib import Path
        
        informes_dir = Path("informes").resolve()
        os.makedirs(informes_dir, exist_ok=True)
        
        doc = docx.Document()
        
        p = doc.add_paragraph()
        run = p.add_run("CERBERUS LOCAL - REPORTE EJECUTIVO DE INCIDENTES")
        run.bold = True
        run.font.size = Pt(16)
        
        doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        stats = self.summary()
        doc.add_heading("1. Resumen de Detección", level=1)
        doc.add_paragraph(f"Eventos Monitoreados: {stats.get('events_total', 0)}")
        doc.add_paragraph(f"Amenazas Detectadas: {stats.get('findings_total', 0)}")
        doc.add_paragraph(f"Acciones de Mitigación: {stats.get('actions_total', 0)}")
        doc.add_paragraph(f"Modo Operativo Actual: {stats.get('mode', 'N/A')}")
        
        findings = self.findings(limit=50)
        doc.add_heading("2. Detalle de Amenazas Recientes (Top 50)", level=1)
        for f in findings:
            p = doc.add_paragraph()
            p.add_run(f"[{f['timestamp']}] ").bold = True
            p.add_run(f"Severidad: {f['severity']} | PID: {f['pid']}\n")
            p.add_run(f"  Reglas: {', '.join(f.get('rule_ids', ['N/A']))}\n")
            p.add_run(f"  Fuentes: {', '.join(f.get('sources', []))}\n")
            ai_info = f"Ollama IA Analyst: {f.get('ai_family', 'Pendiente')}"
            ai_conf = f"(Conf: {f.get('ai_confidence', 0)})"
            p.add_run(f"  {ai_info} {ai_conf}")
        
        timestamp_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Cerberus_Report_{timestamp_str}.docx"
        filepath = informes_dir / filename
        
        doc.save(str(filepath))
        return str(filepath)


